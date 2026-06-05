import re
import os
from docx import Document
from docx.shared import Pt, Cm, Mm, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

ROOT = r'/mnt/c/Users/skinix/cursach'
MD_FILE = os.path.join(ROOT, 'пояснительная_записка', 'Пояснительная_записка.md')
OUT_FILE = os.path.join(ROOT, 'Пояснительная_записка.docx')

IMG_ARCH = os.path.join(ROOT, 'screenshot_architecture.png')
IMG_PANEL = os.path.join(ROOT, 'image.png')
IMG_CLASS = os.path.join(ROOT, 'class_diagram.png')
IMG_COMP = os.path.join(ROOT, 'diagram_0.png')
IMG_SEQ = os.path.join(ROOT, 'diagram_1.png')


def set_cell_font(cell, name='Times New Roman', size=14):
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.name = name
            r.font.size = Pt(size)


def set_default_style(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)
    pf = style.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.first_line_indent = Cm(1.25)

    for level in range(1, 4):
        hs = doc.styles[f'Heading {level}']
        hs.font.name = 'Times New Roman'
        hs.font.size = Pt(14)
        hs.font.bold = True
        hs.font.color.rgb = None
        hs.paragraph_format.first_line_indent = Cm(0)
        hs.paragraph_format.space_before = Pt(6)
        hs.paragraph_format.space_after = Pt(3)


def add_page_number(doc):
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._r.append(fld_char_begin)
        run2 = p.add_run()
        instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2._r.append(instr)
        run3 = p.add_run()
        fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run3._r.append(fld_char_end)
        for r in p.runs:
            r.font.name = 'Times New Roman'
            r.font.size = Pt(14)


def set_margins(doc):
    for section in doc.sections:
        section.left_margin = Mm(30)
        section.right_margin = Mm(10)
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(20)


def add_title_page(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ РФ\n'
                     'ФГБОУ ВО «ТУЛЬСКИЙ ГОСУДАРСТВЕННЫЙ УНИВЕРСИТЕТ»')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Кафедра информационной безопасности')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)

    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('КУРСОВАЯ РАБОТА (ПРОЕКТ)\n'
                     'по дисциплине «Методы и технологии программирования»')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('«3D-редактор для параметрического размещения объектов»\n'
                     'Вариант 46')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True

    for _ in range(4):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run('Выполнил: студент гр. 220032-11\n'
                     'Таминдаров Николай Евгеньевич\n'
                     '\n'
                     'Проверил: ______________\n'
                     '_______________ (___________________)\n'
                     '\n'
                     'Работа защищена с оценке ___________')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Тула 2026')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)

    doc.add_page_break()


def add_toc(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('СОДЕРЖАНИЕ')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True
    doc.add_paragraph()


def parse_markdown_to_docx():
    with open(MD_FILE, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()
    set_default_style(doc)
    set_margins(doc)

    add_title_page(doc)
    add_toc(doc)

    # Track whether we're inside a code block
    in_code = False
    code_lines = []
    in_list = False

    for line in lines:
        stripped = line.rstrip('\n')

        # Handle code blocks
        if stripped.startswith('```'):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                in_code = False
                text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
                p.paragraph_format.first_line_indent = Cm(0)
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(3)
            continue

        if in_code:
            code_lines.append(stripped)
            continue

        # Skip mermaid diagram blocks, comments, YAML front matter
        if stripped.startswith('---') or stripped.startswith('```mermaid'):
            continue

        # Empty line
        if not stripped:
            if in_list:
                in_list = False
            continue

        # Image references: *[image.png — caption]* or *[место для ...]*
        img_match = re.match(r'\*\[(.+?)\]\*', stripped)
        if img_match:
            ref = img_match.group(1)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Cm(0)
            if 'image.png' in ref and os.path.exists(IMG_PANEL):
                run = p.add_run()
                run.add_picture(IMG_PANEL, width=Inches(5.5))
            elif 'архитектур' in ref.lower() and os.path.exists(IMG_ARCH):
                run = p.add_run()
                run.add_picture(IMG_ARCH, width=Inches(5.5))
            elif 'class_diagram' in ref and os.path.exists(IMG_CLASS):
                run = p.add_run()
                run.add_picture(IMG_CLASS, width=Inches(5.5))
            elif 'diagram_0' in ref and os.path.exists(IMG_COMP):
                run = p.add_run()
                run.add_picture(IMG_COMP, width=Inches(5.5))
            elif 'diagram_1' in ref and os.path.exists(IMG_SEQ):
                run = p.add_run()
                run.add_picture(IMG_SEQ, width=Inches(5.5))
            else:
                run = p.add_run(f'[{ref}]')
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.italic = True
            continue

        # Headers
        if stripped.startswith('## '):
            text = stripped[3:]
            doc.add_heading(text, level=2)
            continue
        if stripped.startswith('### '):
            text = stripped[4:]
            doc.add_heading(text, level=3)
            continue
        if stripped.startswith('# '):
            text = stripped[2:]
            doc.add_heading(text, level=1)
            continue

        # Bold text
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', stripped)

        # Tables
        if '|' in stripped and stripped.startswith('|'):
            if '|---' in stripped or '|:---' in stripped:
                continue
            cells = [c.strip() for c in stripped.split('|') if c.strip() != '']
            if len(cells) > 1:
                # Check if table header or data row
                if not hasattr(parse_markdown_to_docx, '_in_table'):
                    parse_markdown_to_docx._in_table = []
                parse_markdown_to_docx._in_table.append(cells)
            continue

        # Check if we have accumulated table data
        if hasattr(parse_markdown_to_docx, '_in_table') and parse_markdown_to_docx._in_table:
            rows = parse_markdown_to_docx._in_table
            parse_markdown_to_docx._in_table = []
            num_cols = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=num_cols)
            table.style = 'Table Grid'
            for i, row_data in enumerate(rows):
                for j, cell_text in enumerate(row_data):
                    if j < num_cols:
                        cell = table.cell(i, j)
                        cell.text = cell_text
                        set_cell_font(cell)
            doc.add_paragraph()
            continue

        # Lists starting with -
        if stripped.startswith('- '):
            in_list = True
            text = stripped[2:]
            p = doc.add_paragraph(style='List Bullet')
            p.text = ''
            run = p.add_run(text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)

    add_page_number(doc)
    return doc


if __name__ == '__main__':
    doc = parse_markdown_to_docx()
    doc.save(OUT_FILE)
    print(f'DOCX saved: {OUT_FILE}')
